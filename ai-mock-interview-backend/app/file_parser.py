import io
import logging
from typing import Union, BinaryIO
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class FileParser:
    """
    Utility class for parsing different file formats and extracting text.
    Supports: PDF, DOCX, TXT
    """
    
    @staticmethod
    def extract_text_from_pdf(file_content: Union[bytes, BinaryIO]) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content as bytes or file object
            
        Returns:
            Extracted text as string
        """
        try:
            if isinstance(file_content, bytes):
                file_content = io.BytesIO(file_content)
            
            reader = PdfReader(file_content)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: Union[bytes, BinaryIO]) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes or file object
            
        Returns:
            Extracted text as string
        """
        try:
            if isinstance(file_content, bytes):
                file_content = io.BytesIO(file_content)
            
            doc = Document(file_content)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_content: Union[bytes, BinaryIO]) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_content: TXT file content as bytes or file object
            
        Returns:
            Extracted text as string
        """
        try:
            if isinstance(file_content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        return file_content.decode(encoding).strip()
                    except UnicodeDecodeError:
                        continue
                raise ValueError("Unable to decode text file with common encodings")
            else:
                return file_content.read().decode('utf-8').strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise ValueError(f"Failed to parse TXT file: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: Union[bytes, BinaryIO], filename: str) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_content: File content as bytes or file object
            filename: Original filename with extension
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file format is not supported
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return FileParser.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return FileParser.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.txt'):
            return FileParser.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format. Please upload PDF, DOCX, or TXT files.")
